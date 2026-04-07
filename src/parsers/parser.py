import pdfplumber
import docx
from abc import ABC, abstractmethod
from pathlib import Path
from docx.oxml.ns import qn
from src.exceptions import PDFParserImportException, WordParserImportException, UnsupportedFileTypeException, ResumeFileNotFoundException, ResumeTooLongException
from src.constants import FileExtension, MAX_FILE_SIZE, MAX_PAGE_COUNT
from src.logger import get_logger

logger = get_logger()

class FileParser(ABC):
    @abstractmethod
    def parse(self, file_path: str):
        pass

class PDFParser(FileParser):
    def parse(self, file_path: str):
        pages = []
        try:
            with pdfplumber.open(file_path) as pdf:
                if len(pdf.pages) > MAX_PAGE_COUNT:
                    logger.error("PDF exceeds page limit: %d pages (max %d)", len(pdf.pages), MAX_PAGE_COUNT)
                    raise ResumeTooLongException(f"Resume exceeds {MAX_PAGE_COUNT} pages ({len(pdf.pages)} pages found)")
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text.strip())
            return "\n".join(pages)
        except ResumeTooLongException:
            raise
        except Exception as e:
            logger.error("PDF parsing failed: %s", e)
            raise PDFParserImportException(f"PDF parsing failed: {e}") from e
        
class WordParser(FileParser):
    def parse(self, file_path: str):
        try:
            doc = docx.Document(file_path)

            page_breaks = sum(
                1
                for para in doc.paragraphs
                for run in para.runs
                for br in run._element.findall(qn('w:br'))
                if br.get(qn('w:type')) == 'page'
            )
            if page_breaks + 1 > MAX_PAGE_COUNT:
                logger.error("Word doc exceeds page limit: %d pages (max %d)", page_breaks + 1, MAX_PAGE_COUNT)
                raise ResumeTooLongException(f"Resume exceeds {MAX_PAGE_COUNT} pages ({page_breaks + 1} pages found)")

            lines = []
    
            # Standard paragraphs (headings, bullets, plain text)
            for para in doc.paragraphs:
                if para.text.strip():
                    lines.append(para.text.strip())
    
            # Table cells — many resume templates use tables for layout
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text and text not in lines:
                            lines.append(text)
    
            return "\n".join(lines)
        except ResumeTooLongException:
            raise
        except Exception as e:
            logger.error("Word parsing failed: %s", e)
            raise WordParserImportException(f"Word parsing failed: {e}") from e

class ParserFactory:
    _registry = {
        FileExtension.PDF: PDFParser,
        FileExtension.DOCX: WordParser,
    }

    @classmethod
    def get_parser(cls, file_path: str) -> FileParser:
        path = Path(file_path)
        if not path.exists():
            logger.error("File not found: %s", file_path)
            raise ResumeFileNotFoundException(f"File not found: {file_path}")

        if path.stat().st_size > MAX_FILE_SIZE:
            logger.error("File exceeds size limit: %s", file_path)
            raise UnsupportedFileTypeException(f"File exceeds maximum allowed size {MAX_FILE_SIZE / (1024 * 1024)} MB")

        ext = path.suffix.lower()
        parser_class = cls._registry.get(ext)
        if parser_class is None:
            logger.error("Unsupported file type: %s", ext)
            raise UnsupportedFileTypeException(f"Unsupported file type: {ext}")
        return parser_class()