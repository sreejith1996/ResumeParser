"""
Shared pytest fixtures for ResumeParser tests.

PDF fixtures are created in-memory using fpdf2 (imported as fpdf).
DOCX fixtures are created in-memory using python-docx.
Corrupt file fixtures write raw bytes to a temp file.
"""
import pytest
from fpdf import FPDF
import docx


def _build_pdf(path, pages_text):
    """Write a PDF with one page per element in pages_text to path."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=False)
    for text in pages_text:
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.multi_cell(0, 8, text)
    pdf.output(str(path))


def _build_docx(path, paragraphs=None, tables=None):
    """Write a DOCX with the given paragraphs and/or tables to path."""
    doc = docx.Document()
    for para in (paragraphs or []):
        doc.add_paragraph(para)
    for table_data in (tables or []):
        rows = len(table_data)
        cols = len(table_data[0]) if rows else 0
        if rows and cols:
            table = doc.add_table(rows=rows, cols=cols)
            for i, row in enumerate(table_data):
                for j, cell_text in enumerate(row):
                    table.cell(i, j).text = cell_text
    doc.save(str(path))


# ── PDF fixtures ────────────────────────────────────────────────────────────

@pytest.fixture
def single_page_pdf(tmp_path):
    """A valid single-page PDF containing 'Hello World'."""
    path = tmp_path / "single.pdf"
    _build_pdf(path, ["Hello World"])
    return str(path)


@pytest.fixture
def multi_page_pdf(tmp_path):
    """A valid two-page PDF."""
    path = tmp_path / "multi.pdf"
    _build_pdf(path, ["Page one content", "Page two content"])
    return str(path)


@pytest.fixture
def empty_page_pdf(tmp_path):
    """A valid PDF with one page that has no text."""
    pdf = FPDF()
    pdf.add_page()  # blank page — no text added
    path = tmp_path / "empty.pdf"
    pdf.output(str(path))
    return str(path)


@pytest.fixture
def whitespace_pdf(tmp_path):
    """A PDF whose page text is all whitespace (should be stripped away)."""
    path = tmp_path / "whitespace.pdf"
    # pdfplumber will return whitespace from the cell, code strips it
    _build_pdf(path, ["   \n   \t   "])
    return str(path)


@pytest.fixture
def corrupt_pdf(tmp_path):
    """A file with a .pdf extension but invalid content."""
    path = tmp_path / "corrupt.pdf"
    path.write_bytes(b"not a valid pdf file at all")
    return str(path)


# ── DOCX fixtures ───────────────────────────────────────────────────────────

@pytest.fixture
def paragraph_docx(tmp_path):
    """A DOCX with only paragraphs (no tables)."""
    path = tmp_path / "paragraphs.docx"
    _build_docx(path, paragraphs=["First paragraph", "Second paragraph"])
    return str(path)


@pytest.fixture
def table_docx(tmp_path):
    """A DOCX with only a table (no paragraphs)."""
    path = tmp_path / "table.docx"
    _build_docx(path, tables=[[["Cell A1", "Cell A2"], ["Cell B1", "Cell B2"]]])
    return str(path)


@pytest.fixture
def mixed_docx(tmp_path):
    """A DOCX with both paragraphs and a table."""
    path = tmp_path / "mixed.docx"
    _build_docx(
        path,
        paragraphs=["Paragraph line"],
        tables=[[["Table cell"]]],
    )
    return str(path)


@pytest.fixture
def duplicate_docx(tmp_path):
    """A DOCX where paragraph text also appears in a table cell (should not duplicate)."""
    path = tmp_path / "duplicate.docx"
    text = "Shared text"
    _build_docx(path, paragraphs=[text], tables=[[[text]]])
    return str(path)


@pytest.fixture
def empty_docx(tmp_path):
    """A DOCX with no paragraphs and no tables."""
    path = tmp_path / "empty.docx"
    _build_docx(path)
    return str(path)


@pytest.fixture
def corrupt_docx(tmp_path):
    """A file with a .docx extension but invalid content."""
    path = tmp_path / "corrupt.docx"
    path.write_bytes(b"not a valid docx file at all")
    return str(path)


# ── Misc fixtures ────────────────────────────────────────────────────────────

@pytest.fixture
def txt_file(tmp_path):
    """A plain .txt file (unsupported type)."""
    path = tmp_path / "resume.txt"
    path.write_text("some text")
    return str(path)


@pytest.fixture
def png_file(tmp_path):
    """A file with .png extension (unsupported type)."""
    path = tmp_path / "resume.png"
    path.write_bytes(b"\x89PNG\r\n")
    return str(path)


@pytest.fixture
def no_extension_file(tmp_path):
    """A file with no extension (unsupported type)."""
    path = tmp_path / "resume"
    path.write_text("content")
    return str(path)


@pytest.fixture
def upper_pdf_file(tmp_path):
    """A file with uppercase .PDF extension."""
    path = tmp_path / "resume.PDF"
    _build_pdf(path, ["content"])
    return str(path)


@pytest.fixture
def upper_docx_file(tmp_path):
    """A file with uppercase .DOCX extension."""
    doc = docx.Document()
    doc.add_paragraph("content")
    path = tmp_path / "resume.DOCX"
    doc.save(str(path))
    return str(path)
