import os
import pytest
from unittest.mock import patch, MagicMock
from fpdf import FPDF
import docx as docx_lib

from src.resume_extractor import ResumeParserFramework
from src.parsers.parser import PDFParser, WordParser
from src.constants import ExtractionField, NEREntityLabel
from src.models.schema import ResumeData


# ── Helpers ──────────────────────────────────────────────────────────────────

SAMPLE_PDF = os.path.join(os.path.dirname(__file__), "..", "samples", "Resume.pdf")


def _build_pdf_bytes(tmp_path, text, filename="resume.pdf"):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(0, 8, text)
    path = tmp_path / filename
    pdf.output(str(path))
    return str(path)


def _build_docx_bytes(tmp_path, paragraphs=None, table_data=None, filename="resume.docx"):
    doc = docx_lib.Document()
    for p in (paragraphs or []):
        doc.add_paragraph(p)
    if table_data:
        rows, cols = len(table_data), len(table_data[0])
        table = doc.add_table(rows=rows, cols=cols)
        for i, row in enumerate(table_data):
            for j, cell in enumerate(row):
                table.cell(i, j).text = cell
    path = tmp_path / filename
    doc.save(str(path))
    return str(path)


def _mock_ner(name):
    """Return a mock NLP that will find `name` as a PERSON entity."""
    mock_nlp = MagicMock()
    mock_doc = MagicMock()
    ent = MagicMock()
    ent.text = name
    ent.label_ = NEREntityLabel.PERSON
    mock_doc.ents = [ent]
    mock_nlp.return_value = mock_doc
    return mock_nlp



class TestPDFIntegration:
    def test_pdf_returns_resume_data_with_expected_fields(self, tmp_path): 
        resume_text = (
            "Alice Brown\n"
            "alice@example.com  |  555-1234\n"
            "Summary: Experienced Python developer.\n"
            "Skills: Python, Docker, Kubernetes\n"
        )
        pdf_path = _build_pdf_bytes(tmp_path, resume_text)

        with patch("src.strategies.extraction_strategies.spacy.load", return_value=_mock_ner("Alice Brown")), \
             patch("src.strategies.extraction_strategies.genai.Client") as mock_cls:
            mock_client = MagicMock()
            mock_cls.return_value = mock_client
            mock_resp = MagicMock()
            mock_resp.text = '["Python", "Docker", "Kubernetes"]'
            mock_client.models.generate_content.return_value = mock_resp

            framework = ResumeParserFramework()
            result = framework.parse_resume(pdf_path)

        assert isinstance(result, ResumeData)
        assert result.name == "Alice Brown"
        assert result.email == "alice@example.com"
        assert "Python" in result.skills

    def test_pdf_no_email_returns_empty_email(self, tmp_path):  
        resume_text = "Bob Lee\nSoftware Engineer\nPython, Go"
        pdf_path = _build_pdf_bytes(tmp_path, resume_text)

        with patch("src.strategies.extraction_strategies.spacy.load", return_value=_mock_ner("Bob Lee")), \
             patch("src.strategies.extraction_strategies.genai.Client") as mock_cls:
            mock_client = MagicMock()
            mock_cls.return_value = mock_client
            mock_resp = MagicMock()
            mock_resp.text = '["Python", "Go"]'
            mock_client.models.generate_content.return_value = mock_resp

            framework = ResumeParserFramework()
            result = framework.parse_resume(pdf_path)

        assert result.email == ""

    def test_pdf_multiple_emails_returns_first(self, tmp_path):  
        resume_text = "Alice Brown\nalice@primary.com\nalice@secondary.com\nPython"
        pdf_path = _build_pdf_bytes(tmp_path, resume_text)

        with patch("src.strategies.extraction_strategies.spacy.load", return_value=_mock_ner("Alice Brown")), \
             patch("src.strategies.extraction_strategies.genai.Client") as mock_cls:
            mock_client = MagicMock()
            mock_cls.return_value = mock_client
            mock_resp = MagicMock()
            mock_resp.text = '["Python"]'
            mock_client.models.generate_content.return_value = mock_resp

            framework = ResumeParserFramework()
            result = framework.parse_resume(pdf_path)

        assert result.email == "alice@primary.com"

    def test_pdf_no_skills_returns_empty_list(self, tmp_path): 
        resume_text = "Alice Brown\nalice@example.com\nSome work experience."
        pdf_path = _build_pdf_bytes(tmp_path, resume_text)

        with patch("src.strategies.extraction_strategies.spacy.load", return_value=_mock_ner("Alice Brown")), \
             patch("src.strategies.extraction_strategies.genai.Client") as mock_cls:
            mock_client = MagicMock()
            mock_cls.return_value = mock_client
            mock_resp = MagicMock()
            mock_resp.text = "[]"
            mock_client.models.generate_content.return_value = mock_resp

            framework = ResumeParserFramework()
            result = framework.parse_resume(pdf_path)

        assert result.skills == []

    def test_pdf_no_name_returns_empty_string(self, tmp_path):  
        # NER and fallback both fail to find a name
        resume_text = "123 Main St.\nCity, State 12345\nalice@example.com"
        pdf_path = _build_pdf_bytes(tmp_path, resume_text)

        mock_nlp = MagicMock()
        mock_doc = MagicMock()
        mock_doc.ents = []
        mock_nlp.return_value = mock_doc

        with patch("src.strategies.extraction_strategies.spacy.load", return_value=mock_nlp), \
             patch("src.strategies.extraction_strategies.genai.Client") as mock_cls:
            mock_client = MagicMock()
            mock_cls.return_value = mock_client
            mock_resp = MagicMock()
            mock_resp.text = "[]"
            mock_client.models.generate_content.return_value = mock_resp

            framework = ResumeParserFramework()
            result = framework.parse_resume(pdf_path)

        assert result.name == ""

    def test_pdf_all_caps_header_name_extracted(self, tmp_path): 
        resume_text = "ALICE BROWN\nALICE@EXAMPLE.COM\nPYTHON DEVELOPER"
        pdf_path = _build_pdf_bytes(tmp_path, resume_text)

        with patch("src.strategies.extraction_strategies.spacy.load", return_value=_mock_ner("Alice Brown")), \
             patch("src.strategies.extraction_strategies.genai.Client") as mock_cls:
            mock_client = MagicMock()
            mock_cls.return_value = mock_client
            mock_resp = MagicMock()
            mock_resp.text = '["Python"]'
            mock_client.models.generate_content.return_value = mock_resp

            framework = ResumeParserFramework()
            result = framework.parse_resume(pdf_path)

        assert result.name == "Alice Brown"



class TestDOCXIntegration:
    def test_docx_returns_resume_data(self, tmp_path):    
        paragraphs = [
            "Alice Brown",
            "alice@example.com",
            "Python Developer",
        ]
        docx_path = _build_docx_bytes(tmp_path, paragraphs=paragraphs)

        with patch("src.strategies.extraction_strategies.spacy.load", return_value=_mock_ner("Alice Brown")), \
             patch("src.strategies.extraction_strategies.genai.Client") as mock_cls:
            mock_client = MagicMock()
            mock_cls.return_value = mock_client
            mock_resp = MagicMock()
            mock_resp.text = '["Python"]'
            mock_client.models.generate_content.return_value = mock_resp

            framework = ResumeParserFramework()
            result = framework.parse_resume(docx_path)

        assert isinstance(result, ResumeData)
        assert result.name == "Alice Brown"
        assert result.email == "alice@example.com"
        assert result.skills == ["Python"]

    def test_docx_table_layout_extracts_correctly(self, tmp_path): 
        # Many DOCX resumes use a table for layout
        table_data = [
            ["Alice Brown", "alice@example.com"],
            ["Python Developer", "5 years experience"],
        ]
        docx_path = _build_docx_bytes(tmp_path, table_data=table_data)

        with patch("src.strategies.extraction_strategies.spacy.load", return_value=_mock_ner("Alice Brown")), \
             patch("src.strategies.extraction_strategies.genai.Client") as mock_cls:
            mock_client = MagicMock()
            mock_cls.return_value = mock_client
            mock_resp = MagicMock()
            mock_resp.text = '["Python"]'
            mock_client.models.generate_content.return_value = mock_resp

            framework = ResumeParserFramework()
            result = framework.parse_resume(docx_path)

        assert result.name == "Alice Brown"
        assert result.email == "alice@example.com"


# ── I-09: main.py smoke test ──────────────────────────────────────────────────

class TestMainSmoke:
    def test_main_runs_without_error(self):
        """main() should complete without raising when all dependencies are mocked."""
        import main as main_module

        mock_data = ResumeData(name="Alice Brown", email="alice@example.com", skills=["Python"])

        with patch.object(main_module, "load_dotenv"), \
             patch.object(main_module, "ResumeParserFramework") as mock_fw_cls:
            mock_fw = MagicMock()
            mock_fw.parse_resume.return_value = mock_data
            mock_fw_cls.return_value = mock_fw
            main_module.main()


# ── Live tests (require real API key + model) ─────────────────────────────────

@pytest.mark.live
@pytest.mark.skipif(
    not os.path.exists(SAMPLE_PDF),
    reason="samples/Resume.pdf not found",
)
class TestLiveIntegration:
    """
    These tests hit the real Gemini API and load the real spaCy model.
    Run with: pytest -m live
    """

    def test_sample_pdf_returns_non_empty_name(self):
        framework = ResumeParserFramework()
        result = framework.parse_resume(SAMPLE_PDF)
        assert result.name != ""

    def test_sample_pdf_returns_valid_email(self):
        framework = ResumeParserFramework()
        result = framework.parse_resume(SAMPLE_PDF)
        assert "@" in result.email

    def test_sample_pdf_returns_non_empty_skills(self):
        framework = ResumeParserFramework()
        result = framework.parse_resume(SAMPLE_PDF)
        assert len(result.skills) > 0
